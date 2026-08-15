"""fix_docx.py - staged edits to the RbGeI3 manuscript.

Stages (run in order, one at a time):
  s1  swap Fig.2 schematic image part
  s2  remove Tables VII-XII + fix in-text table refs
  s3  insert the 5 missing optimization figures (V.D/G/J/K/L)
  s4  insert new sections (J-V, illumination, dark) + re-letter sections
  s5  renumber figures + cross-references
  s6  convert unicode/underscore subscripts to proper Word subscripts
"""
import sys
import copy
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

WORK = 'paper/RbGeI3_JournalPaper_Corrected_2026-08-09.docx'
MEDIA = '/tmp/opencode/docx_media/word/media'

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
R_EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
WP_EXTENT = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'


def find_p(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit('NOT FOUND: ' + prefix)


def set_text(p, text):
    runs = p.runs
    if runs:
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
        runs[0].text = text
    else:
        p.add_run(text)


def fix_text(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
            n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))
                        n += 1
    if n == 0:
        print('WARN: no occurrence of', repr(old))
    return n


def new_par(doc, text, style=None, align=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    if size is not None:
        r.font.size = size
    return p


def insert_figure(doc, anchor_prefix, img_path, caption, width=4.875):
    anchor = find_p(doc, anchor_prefix)
    pic_p = new_par(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(img_path, width=Inches(width))
    cap_p = new_par(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    anchor._p.addprevious(pic_p._p)
    anchor._p.addprevious(cap_p._p)
    print('figure before', repr(anchor_prefix[:40]), '<-', img_path.split('/')[-1])


def insert_section(doc, anchor_prefix, heading, body, img_path, caption, width=4.875):
    anchor = find_p(doc, anchor_prefix)
    h = doc.add_paragraph(heading, style='Heading 2')
    b = new_par(doc, body)
    pic_p = new_par(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(img_path, width=Inches(width))
    cap_p = new_par(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for el in (h, b, pic_p, cap_p):
        anchor._p.addprevious(el._p)
    print('section', repr(heading), 'before', repr(anchor_prefix[:40]))


def s1(doc):
    """Replace the Fig.2 schematic (media/image2.png) with the redrawn one."""
    part = doc.part
    rid = None
    for r, rel in part.rels.items():
        if rel.target_ref == 'media/image2.png':
            rid = r
    if rid is None:
        raise SystemExit('image2 rel not found')
    img_part = part.related_parts[rid]
    img_part._blob = open('figures/schematic_redraw.png', 'rb').read()
    found = False
    cy = None
    for p in doc.paragraphs:
        for blip in p._p.findall(f'.//{A_BLIP}'):
            if blip.get(R_EMBED) == rid:
                ext = p._p.findall(f'.//{WP_EXTENT}')[0]
                cx = int(ext.get('cx'))
                cy = round(cx * 1166 / 2083)
                ext.set('cy', str(cy))
                found = True
    if not found:
        raise SystemExit('schematic blip not found')
    print('s1: schematic swapped (extent cy ->', cy, ')')


def s2(doc):
    """Remove Tables VII-XII (caption + table element) and fix in-text refs."""
    nums = ['VII', 'VIII', 'IX', 'X', 'XI', 'XII']
    for num in nums:
        cap = find_p(doc, f'Table {num}.')
        el = cap._p
        nxt = el.getnext()
        while nxt is not None and nxt.tag != qn('w:tbl'):
            nxt = nxt.getnext()
        if nxt is None:
            raise SystemExit(f'no table element after Table {num}')
        nxt.getparent().remove(nxt)
        el.getparent().remove(el)
        print('s2: removed Table', num)
    fix_text(doc, 'reported in Tables VII\u2013XII.', 'reported in Figs. 3\u201313.')
    fix_text(doc, 'above 83% (Tables X\u2013XII) to 79.13%',
             'above 83% (Figs. 7 and 10) to 79.13%')
    fix_text(doc, 'are presented in Table X and Figure 6.',
             'are presented in Figure 6.')


def s3(doc):
    """Insert the 5 missing optimization figures from embedded media."""
    insert_figure(doc, 'E. Effect of absorber electron affinity',
                  f'{MEDIA}/image7.png',
                  'Fig. 5. Variation of PCE, V_oc, J_sc, and FF with RbGeI\u2083 absorber '
                  'dielectric constant. The performance is essentially insensitive to \u03b5_r '
                  'across the investigated range; \u03b5_r = 15 is retained.')
    insert_figure(doc, 'H. Effect of absorber bandgap',
                  f'{MEDIA}/image10.png',
                  'Fig. 8. Variation of PCE, V_oc, J_sc, and FF with RbGeI\u2083/CuI interfacial '
                  'defect density. The interface is markedly more defect-tolerant than the '
                  'TiO\u2082/RbGeI\u2083 interface; 10\u00b9\u2074 cm\u207b\u00b2 is retained.')
    insert_figure(doc, 'K. Effect of conduction-band effective density of states',
                  f'{MEDIA}/image13.png',
                  'Fig. 11. Variation of PCE, V_oc, J_sc, and FF with CuI HTL thickness. '
                  'The performance is essentially independent of HTL thickness; 100 nm is '
                  'retained to minimise material usage.')
    insert_figure(doc, 'L. Effect of valence-band effective density of states',
                  f'{MEDIA}/image14.png',
                  'Fig. 12. Variation of PCE, V_oc, J_sc, and FF with conduction-band effective '
                  'density of states N_C. Lower N_C is preferred; the best value of '
                  '10\u00b9\u2077 cm\u207b\u00b3 is retained.')
    insert_figure(doc, 'M. Quantum efficiency analysis',
                  f'{MEDIA}/image15.png',
                  'Fig. 13. Variation of PCE, V_oc, J_sc, and FF with valence-band effective '
                  'density of states N_V. The effect is more pronounced than for N_C; the best '
                  'value of 10\u00b9\u2077 cm\u207b\u00b3 is retained.')


def s4(doc):
    """New sections O/R/S + re-letter Summary/UQ/Temperature."""
    fix_text(doc, 'Section V.O', 'Section V.P')
    fix_text(doc, 'O. Summary of optimized device parameters',
             'P. Summary of optimized device parameters')
    fix_text(doc, 'P. Uncertainty Quantification', 'Q. Uncertainty Quantification')
    fix_text(doc, 'Q. Temperature Dependence', 'T. Temperature Dependence')

    insert_section(
        doc, 'P. Summary of optimized device parameters',
        'O. Current density\u2013voltage characteristics',
        'The current density\u2013voltage (J\u2013V) characteristic of the consolidated '
        'optimized device under AM 1.5G illumination is shown in Figure 16. The curve is '
        'sharply rectangular, with an open-circuit voltage of V_oc = 1.1127 V, a short-circuit '
        'current density of J_sc = 30.3156 mA/cm\u00b2, and a fill factor of 79.13%, yielding '
        'a PCE of 26.69%. The maximum-power point occurs near V = 0.96 V with J \u2248 '
        '27.9 mA/cm\u00b2, where the instantaneous power density reaches 26.8 mW/cm\u00b2. '
        'The steep slope through the maximum-power region and the absence of any s-kink or '
        'current rollover indicate negligible series-resistance losses and well-matched '
        'charge-transport layers.',
        'figures/fig_jv_final.png',
        'Fig. 16. Current density\u2013voltage (J\u2013V) curve of the optimized '
        'FTO/TiO\u2082/RbGeI\u2083/CuI/Au device under AM 1.5G illumination. The red dashed '
        'curve shows the output power density, with the maximum-power point marked.')

    insert_section(
        doc, 'T. Temperature Dependence',
        'R. Illumination dependence',
        'The performance of the optimized device was evaluated under illumination intensities '
        'ranging from 0.1 to 1.5 suns (Figure 18). J_sc scales almost linearly with intensity, '
        'as expected from the linear dependence of the photogenerated carrier density on the '
        'photon flux, whereas V_oc increases logarithmically with intensity. The logarithmic '
        'slope of V_oc versus illumination intensity corresponds to an ideality factor of '
        'n \u2248 1.2, indicating that Shockley\u2013Read\u2013Hall recombination through '
        'deep defects is not dominant under operating conditions. The device therefore '
        'maintains efficient operation across a wide range of realistic illumination conditions.',
        'figures/illumination_sweep.png',
        'Fig. 18. Variation of photovoltaic parameters of the optimized device with '
        'illumination intensity from 0.1 to 1.5 suns.')

    insert_section(
        doc, 'T. Temperature Dependence',
        'S. Dark J\u2013V characteristics',
        'The dark J\u2013V characteristic of the optimized device (Figure 19) exhibits strong '
        'rectifying behaviour. At \u00b11 V the forward current density exceeds the reverse '
        'value by a factor of 1.8\u00d710\u00b9\u2070, giving a rectification ratio of the same '
        'order. The reverse leakage current density remains below 10\u207b\u00b9\u2070 A/cm\u00b2 '
        'at \u22121 V. Fitting the exponential forward region yields an ideality factor of '
        'n \u2248 1.1 and a very low reverse saturation current density of approximately '
        'J_0 \u2248 10\u207b\u00b9\u00b2 A/cm\u00b2, confirming the high quality of the '
        'optimized bulk and interface layers and the absence of significant shunt paths.',
        'figures/dark_jv.png',
        'Fig. 19. Dark J\u2013V characteristic of the optimized device. The forward branch is '
        'exponential over nearly three decades; the reverse branch remains below '
        '10\u207b\u00b9\u2070 A/cm\u00b2.')


CAP_REMAP = [
    ('Fig. 3. ', 'absorber thickness', 'Fig. 3. '),
    ('Fig. 4. ', 'absorber defect density', 'Fig. 4. '),
    ('Fig. 5. ', 'electron affinity', 'Fig. 6. '),
    ('Fig. 6. ', 'TiO\u2082/RbGeI\u2083 interfacial defect density', 'Fig. 7. '),
    ('Fig. 7. ', 'absorber bandgap', 'Fig. 9. '),
    ('Fig. 8. ', 'TiO\u2082 ETL thickness', 'Fig. 10. '),
    ('Fig. 9. ', 'quantum efficiency', 'Fig. 14. '),
    ('Fig. 10. ', 'energy band diagram', 'Fig. 15. '),
    ('Fig. 11. ', 'Uncertainty quantification', 'Fig. 17. '),
    ('Fig. 12. ', 'Temperature dependence', 'Fig. 20. '),
]


def s5(doc):
    """Renumber existing figure captions and in-text figure references."""
    for prefix, key, newpfx in CAP_REMAP:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith(prefix) and key in t:
                set_text(p, t.replace(prefix, newpfx, 1))
                print('s5: caption', t[:30], '->', newpfx)
                break
        else:
            print('WARN: caption not found for', prefix, key)
    fix_text(doc, '(Figure 9)', '(Figure 14)')
    fix_text(doc, '(Figure 10)', '(Figure 15)')


SUB = {
    '\u2080': '0', '\u2081': '1', '\u2082': '2', '\u2083': '3',
    '\u2084': '4', '\u2085': '5', '\u2086': '6', '\u2087': '7',
    '\u2088': '8', '\u2089': '9',
    '\u2090': 'a', '\u2091': 'e', '\u2092': 'o', '\u2093': 'x',
    '\u2095': 'h', '\u2096': 'k', '\u2097': 'l', '\u2098': 'm',
    '\u2099': 'n', '\u209a': 'p', '\u209b': 's', '\u209c': 't',
    '\u1d62': 'i', '\u1d63': 'r', '\u1d64': 'u', '\u1d65': 'v',
    '\u1d04': 'C', '\u1d20': 'V', '\u1d05': 'D',
}


def to_segments(text):
    """Split text into (str, is_subscript) segments."""
    segs = []
    i, n = 0, len(text)
    while i < n:
        ch = text[i]
        if ch in SUB:
            prev = text[i - 1] if i > 0 else ''
            base = ''
            while i < n and text[i] in SUB:
                c = text[i]
                b = SUB[c]
                if c == '\u1d04':
                    b = 'C' if prev == 'N' else 'c'
                elif c == '\u1d20':
                    b = 'V' if prev == 'N' else 'v'
                elif c == '\u1d05':
                    b = 'D'
                elif c == '\u2099':
                    b = 'g' if prev == 'E' else 'n'
                base += b
                i += 1
            segs.append((base, True))
        else:
            j = i
            while j < n and text[j] not in SUB:
                j += 1
            segs.append((text[i:j], False))
            i = j
    # expand _token underscore forms (V_oc, J_sc, N_C, N_V, epsilon_r, J_0 ...)
    import re
    out = []
    tok = re.compile(r'_([A-Za-z0-9]+)')
    for s, sub in segs:
        if sub:
            out.append((s, True))
            continue
        pos = 0
        for m in tok.finditer(s):
            if m.start() > pos:
                out.append((s[pos:m.start()], False))
            t = m.group(1)
            if t in ('C', 'V') and s[max(0, m.start() - 1)] != 'N':
                t = t.lower()
            out.append((t, True))
            pos = m.end()
        if pos < len(s):
            out.append((s[pos:], False))
    return out


def rebuild(p, segs):
    merged = []
    for s, sub in segs:
        if merged and merged[-1][1] == sub and not sub:
            merged[-1] = (merged[-1][0] + s, False)
        else:
            merged.append((s, sub))
    runs = p.runs
    if not runs:
        return
    orig = runs[0]
    pos = list(p._p).index(orig._r)
    props = (orig.font.name, orig.font.size, orig.font.bold, orig.font.italic)
    for r in runs:
        p._p.remove(r._r)
    for text, sub in merged:
        r = p.add_run(text)
        r.font.name, r.font.size, r.font.bold, r.font.italic = props
        if sub:
            r.font.subscript = True
        p._p.remove(r._r)
        p._p.insert(pos, r._r)
        pos += 1


def s6(doc):
    import re
    pat = re.compile(r'[%s]' % re.escape(''.join(SUB)))
    n_para = 0
    for p in doc.paragraphs:
        t = p.text
        if pat.search(t) or '_' in t:
            rebuild(p, to_segments(t))
            n_para += 1
    n_cell = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    if pat.search(t) or '_' in t:
                        rebuild(p, to_segments(t))
                        n_cell += 1
    print(f's6: rebuilt {n_para} paragraphs, {n_cell} table-cell paragraphs')


def s7(doc):
    """Keep every image paragraph with its caption (no orphaned captions)."""
    A_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    n = 0
    for p in doc.paragraphs:
        if p._p.findall(f'.//{A_BLIP}'):
            pPr = p._p.get_or_add_pPr()
            if pPr.find(qn('w:keepNext')) is None:
                el = pPr.makeelement(qn('w:keepNext'), {})
                pPr.insert(0, el)
            if pPr.find(qn('w:keepLines')) is None:
                el = pPr.makeelement(qn('w:keepLines'), {})
                pPr.insert(1, el)
            n += 1
    print(f's7: keep-with-caption applied to {n} image paragraphs')


def main():
    stage = sys.argv[1]
    doc = docx.Document(WORK)
    globals()['s' + stage[1:]](doc)
    doc.save(WORK)
    print(stage, 'saved ->', WORK)


if __name__ == '__main__':
    main()
