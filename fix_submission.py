"""fix_submission.py - apply the submission-audit fixes to the manuscript.

Run-level text edits (preserving subscript formatting):
  [17,18,13],,            -> [13,17,18].
  presented in Figure 6.  -> presented in Figure 7.   (Sec F is Fig 7)
  summarised              -> summarized              (2x, consistency)
  n \u2248 1.2            -> n \u2248 1.17           (matches new Fig 18)
  -1.08 mV/K, -0.0313 %/K -> -1.10 mV/K, -0.0304 %/K (matches new Fig 20)
  (80.53 %)               -> (80.27 %)               (matches new Fig 20)
  \u2148 (broken glyph)   -> c                       (J_s + c, V_o + c)

Image swaps (by media part):
  image1  -> figures/gui_screenshot.png     (Fig. 1, cropped GUI)
  image11 -> figures/bandgap_sweep.png      (Fig. 9, proper bandgap sweep)
  image24 -> figures/illumination_sweep.png (Fig. 18, n \u2248 1.17 annotation)
  image4  -> figures/temperature_sweep.png  (Fig. 20, corrected slopes)

Table shading: strip light-blue D9E2F3 cell shading from all tables.
"""
import docx
from docx.oxml.ns import qn

WORK = 'paper/RbGeI3_JournalPaper_Corrected_2026-08-09.docx'
A_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
R_EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
WP_EXTENT = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'

SWAPS = {
    'image1.png': 'figures/gui_screenshot.png',
    'image11.png': 'figures/bandgap_sweep.png',
    'image24.png': 'figures/illumination_sweep.png',
    'image4.png': 'figures/temperature_sweep.png',
}

TEXT_FIXES = [
    ('[17,18,13],,', '[13,17,18].'),
    ('presented in Figure 6. Device performance', 'presented in Figure 7. Device performance'),
    ('summarised', 'summarized'),
    ('ideality factor of n \u2248 1.2,', 'ideality factor of n \u2248 1.17,'),
    ('\u22121.08 mV/K, dPCE/dT = \u22120.0313',
     '\u22121.10 mV/K, dPCE/dT = \u22120.0304'),
    ('FF peaks at 340 K (80.53 %)', 'FF peaks at 340 K (80.27 %)'),
]


def fix_runs(p, old, new):
    n = 0
    for r in p.findall(qn('w:r')):
        ts = r.findall(qn('w:t'))
        for t in ts:
            if t.text and old in t.text:
                t.text = t.text.replace(old, new)
                n += 1
    return n


def fix_broken_glyph(p):
    n = 0
    for t in p.findall(qn('w:r') + '/' + qn('w:t')):
        if t.text and '\u2148' in t.text:
            t.text = t.text.replace('\u2148', 'c')
            n += 1
    return n


def swap_image(doc, media_name, new_path, fix_ratio=False):
    part = doc.part
    rid = None
    for r, rel in part.rels.items():
        if rel.target_ref == f'media/{media_name}':
            rid = r
    if rid is None:
        raise SystemExit(f'rel for {media_name} not found')
    img_part = part.related_parts[rid]
    data = open(new_path, 'rb').read()
    img_part._blob = data
    if fix_ratio:
        from PIL import Image
        import io
        w, h = Image.open(io.BytesIO(data)).size
        for p in doc.paragraphs:
            for blip in p._p.findall(f'.//{A_BLIP}'):
                if blip.get(R_EMBED) == rid:
                    ext = p._p.findall(f'.//{WP_EXTENT}')[0]
                    cx = int(ext.get('cx'))
                    ext.set('cy', str(round(cx * h / w)))
    print(f'swap {media_name} <- {new_path}')


def strip_table_shading(doc):
    n = 0
    for tbl in doc.tables:
        for tcPr in tbl._tbl.iter(qn('w:tcPr')):
            for shd in tcPr.findall(qn('w:shd')):
                fill = shd.get(qn('w:fill'))
                if fill and fill != 'auto':
                    tcPr.remove(shd)
                    n += 1
    print(f'strip_table_shading: removed {n} cell shading fills')


def main():
    doc = docx.Document(WORK)

    # --- text fixes (run-level, preserves subscripts) ---
    for old, new in TEXT_FIXES:
        n = 0
        for p in doc.paragraphs:
            n += fix_runs(p._p, old, new)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        n += fix_runs(p._p, old, new)
        print(f'fix {old!r} -> {new!r}: {n} run(s)' if n else f'WARN: no match for {old!r}')

    # --- broken glyph \u2148 -> c ---
    n = 0
    for p in doc.paragraphs:
        n += fix_broken_glyph(p._p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += fix_broken_glyph(p._p)
    print(f'fix broken glyph: {n} run(s)')

    # --- image swaps ---
    for media, path in SWAPS.items():
        swap_image(doc, media, path, fix_ratio=(media == 'image1.png'))

    # --- table shading ---
    strip_table_shading(doc)

    doc.save(WORK)
    print('saved ->', WORK)


if __name__ == '__main__':
    main()
