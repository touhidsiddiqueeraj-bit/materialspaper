#!/usr/bin/env python3
"""Consistency checker for RbGeI3_Round3.docx/.pdf vs raw data sources."""
import json, math, re, subprocess, sys
import numpy as np
import docx

ROOT = '/home/touhid/Documents/materilaspaper'
DOCX = f'{ROOT}/paper/round3/RbGeI3_Round3.docx'
PDF = f'{ROOT}/paper/round3/RbGeI3_Round3.pdf'
THESIS = '/tmp/opencode/round3/thesis_tables.json'
RAW = f'{ROOT}/fix_results/raw_results.json'
CONV = f'{ROOT}/dark_and_conv_results.json'
UQ = f'{ROOT}/uq_and_dark_results.json'

fails, passes = [], []

def check(name, ok, detail=''):
    (passes if ok else fails).append((name, detail))

def page_count():
    return int(subprocess.run(['pdfinfo', PDF], capture_output=True, text=True)
               .stdout.split('Pages:')[1].split()[0])

def pdf_text():
    return subprocess.run(['pdftotext', '-layout', PDF, '-'], capture_output=True, text=True).stdout

def approx(a, b, tol):
    return abs(a - b) <= tol

text = pdf_text()
text_flat = re.sub(r'\s+', ' ', text)
doc = docx.Document(DOCX)
paras = [p.text for p in doc.paragraphs]
join_all = '\n'.join(paras)
tbl = json.load(open(THESIS))
raw = json.load(open(RAW))
conv = json.load(open(CONV))
uq = json.load(open(UQ))

# ---- A. headline device ----
v, j, ff, pce = 1.12, 30.32, 78.92, 26.80
check('A1  PCE = Voc*Jsc*FF/100', approx(v*j*ff/100, pce, 0.02), f'{v*j*ff/100:.3f} vs {pce}')
fine = conv['FINE']
check('A2  headline traces to FINE run', abs(fine['PCE']-26.80) < 0.02 and approx(fine['FF'], 78.92, 0.01),
      f"FINE PCE={fine['PCE']:.4f} FF={fine['FF']:.4f}")
check('A3  Voc 1.12 in text', 'Voc = 1.12 V' in text)
check('A4  max-power V=0.96 J=27.9 -> 26.8 mW/cm2', approx(0.96*27.9, 26.8, 0.3))

# ---- B. sweep claims vs thesis tables ----
def row(tk, key):
    for r in tbl[tk]['rows']:
        if r[0] == key:
            return [float(x) for x in r[1:]]
    raise KeyError((tk, key))

P, Vc, Jc, F = 0, 1, 2, 3
b = [
    ('Fig 3/V.B 0.7um peak', 'T7', '0.7', 'PCE', 20.70, P),
    ('V.B Jsc start', 'T7', '0.3', 'Jsc', 28.14, Jc),
    ('V.B Jsc end', 'T7', '1', 'Jsc', 34.11, Jc),
    ('V.B Voc start', 'T7', '0.3', 'Voc', 0.823, Vc),
    ('V.B Voc end', 'T7', '1', 'Voc', 0.767, Vc),
    ('V.C 1e14 PCE', 'T8', '1014', 'PCE', 20.40, P),
    ('V.C 1e18 PCE', 'T8', '1018', 'PCE', 7.58, P),
    ('V.C Voc 1e14', 'T8', '1014', 'Voc', 0.848, Vc),
    ('V.C Voc 1e18', 'T8', '1018', 'Voc', 0.550, Vc),
    ('V.E eps 15 PCE', 'T9', '15', 'PCE', 19.92, P),
    ('V.E eps 23.1 PCE', 'T9', '23.1', 'PCE', 19.79, P),
    ('V.E eps 15 Voc', 'T9', '15', 'Voc', 0.819, Vc),
    ('V.E eps 23.1 Voc', 'T9', '23.1', 'Voc', 0.812, Vc),
    ('V.F 1e12 PCE', 'T11', '1012', 'PCE', 21.04, P),
    ('V.F 1e18 PCE', 'T11', '1018', 'PCE', 8.40, P),
    ('V.F Voc 1e12', 'T11', '1012', 'Voc', 0.827, Vc),
    ('V.F Voc 1e18', 'T11', '1018', 'Voc', 0.439, Vc),
    ('V.F Jsc 1e12', 'T11', '1012', 'Jsc', 30.46, Jc),
    ('V.F Jsc 1e18', 'T11', '1018', 'Jsc', 28.09, Jc),
    ('V.F FF 1e12', 'T11', '1012', 'FF', 83.52, F),
    ('V.F FF 1e18', 'T11', '1018', 'FF', 68.11, F),
    ('V.G 1e12 PCE', 'T12', '1012', 'PCE', 19.79, P),
    ('V.G 1e18 PCE', 'T12', '1018', 'PCE', 18.19, P),
    ('V.I 10nm PCE', 'T14', '0.01', 'PCE', 21.12, P),
    ('V.I 0.04 PCE', 'T14', '0.04', 'PCE', 18.72, P),
    ('V.I 0.09 PCE', 'T14', '0.09', 'PCE', 19.97, P),
    ('V.K 1e17 PCE', 'T16', '1×1017', 'PCE', 20.65, P),
    ('V.K 1e20 PCE', 'T16', '1×1020', 'PCE', 18.15, P),
    ('V.K Voc 1e17', 'T16', '1×1017', 'Voc', 0.859, Vc),
    ('V.K Voc 1e20', 'T16', '1×1020', 'Voc', 0.764, Vc),
    ('V.L 1e17 PCE', 'T17', '1×1017', 'PCE', 23.54, P),
    ('V.L 1e20 PCE', 'T17', '1×1020', 'PCE', 18.86, P),
    ('V.L Voc 1e17', 'T17', '1×1017', 'Voc', 0.945, Vc),
    ('V.L Voc 1e20', 'T17', '1×1020', 'Voc', 0.780, Vc),
]
for name, tk, key, m, q, col in b:
    v_ = row(tk, key)[col]
    check(f'{name} ({m})', approx(v_, q, 0.02 if col in (P, F) else 0.012), f'raw {v_:.4f} vs quoted {q}')
t15 = [float(x[1]) for x in tbl['T15']['rows']]
check('V.J CuI PCE spread < 0.002 pp', max(t15)-min(t15) < 0.002, f'span {max(t15)-min(t15):.5f}')
vk = row('T16', '1×1017')[Jc]
check('V.K Jsc ~30.46 constant', approx(vk, 30.4604, 0.002), f'{vk:.4f}')

# ---- C. bandgap sweep vs corrected raw runs ----
for k, m, q in [(('A_Eg1.30'), 'Voc', 0.840), (('A_Eg1.30'), 'Jsc', 33.59),
                (('A_Eg1.70'), 'Voc', 1.002), (('A_Eg1.70'), 'Jsc', 20.30)]:
    r = raw[k]
    check(f'V.H {k} {m} vs corrected run', approx(r[m], q, 0.012), f'raw {r[m]:.4f} vs {q}')

# ---- D. Table VI arithmetic from docx ----
vi = doc.tables[5]
rows = vi.rows[1:]
for r in rows:
    c = [x.text.strip() for x in r.cells]
    voc_, jsc_, ff_, p_ = float(c[2]), float(c[3]), float(c[4]), float(c[5])
    check(f'Table VI {c[0]} arithmetic', approx(voc_*jsc_*ff_/100, p_, 0.03),
          f'{voc_*jsc_*ff_/100:.2f} vs {p_}')
check('Table VI has 8 device rows', len(rows) == 8, str(len(rows)))

# ---- E. UQ statistics ----
pce_uq = np.array([x['PCE'] for x in uq['uq']])
voc_uq = np.array([x['Voc'] for x in uq['uq']])
check('UQ mean 25.92', approx(pce_uq.mean(), 25.92, 0.01), f'{pce_uq.mean():.3f}')
check('UQ sd 1.78', approx(pce_uq.std(ddof=1), 1.78, 0.01), f'{pce_uq.std(ddof=1):.3f}')
check('UQ median 25.85', approx(np.median(pce_uq), 25.85, 0.01), f'{np.median(pce_uq):.3f}')
check('UQ p5 23.38 (linear pct)', approx(np.percentile(pce_uq, 5), 23.38, 0.02), f'{np.percentile(pce_uq,5):.3f}')
check('UQ p95 28.81 (linear pct)', approx(np.percentile(pce_uq, 95), 28.81, 0.02), f'{np.percentile(pce_uq,95):.3f}')
check('UQ Voc mean 1.087', approx(voc_uq.mean(), 1.087, 0.001), f'{voc_uq.mean():.3f}')
check('UQ Voc sd 0.021 in text', 'Mean Voc = 1.087 \u00b1 0.021 V' in text_flat)

# ---- F. dark JV ----
dj = conv['DARK_JV']
V_ = np.array(dj['V']); J_ = np.array(dj['J'])
ip1 = float(np.interp(1.0, V_, J_)); im05 = float(np.interp(-0.5, V_, J_))
check('dark bias window -0.5..+1.3 V', approx(V_.min(), -0.5, 1e-9) and approx(V_.max(), 1.3, 1e-9),
      f'{V_.min():.2f}..{V_.max():.2f}')
check('dark ratio 1.8e10', approx(ip1/abs(im05), 1.8e10, 0.3e10), f'{ip1/abs(im05):.2e}')
check('dark reverse < 1e-10 A/cm2', max(abs(J_[V_ < 0])) < 1e-10, f'{max(abs(J_[V_<0])):.2e}')
check('dark n ~ 1.5 in text', 'ideality factor of n \u2248 1.5' in text_flat)
check('dark J0 ~ 5e-11 in text', 'J0 \u2248 5\u00d710\u207b\u00b9\u00b9 A/cm\u00b2' in text_flat)

# ---- G. bandgap math ----
check('lambda_c 1.4eV = 885.7 nm', approx(1240/1.4, 885.7, 0.1) and '885.7 nm' in text)
check('lambda_c 2.0eV = 620 nm', approx(1240/2.0, 620.0, 0.1) and '620 nm' in text)
check('graded 885.7->1033.3 in seven linear steps',
      'graded from \u2248885.7 nm at the front to \u22481033.3 nm at the back in seven linear steps' in text_flat)

# ---- H. references (from docx, source of truth) ----
refs = [p.text.strip() for p in doc.paragraphs if re.match(r'^\[\d+\]', p.text.strip())]
nums = [int(m.group(1)) for m in (re.match(r'^\[(\d+)\]', r) for r in refs) if m]
check('refs contiguous 1..43 in docx', nums == list(range(1, 44)),
      f'{len(refs)} refs, first {nums[:2]} last {nums[-2:]}')
refs_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith('References'))
body = '\n'.join(p.text for p in doc.paragraphs[:refs_start])
cites = [int(x) for c in re.findall(r'\[([\d,\s\u2013]+)\]', body) for x in re.split(r'[, \u2013]+', c) if x.isdigit()]
firsts, seen = [], set()
for c in cites:
    if c not in seen:
        seen.add(c); firsts.append(c)
ordered = all(firsts[i] <= firsts[i+1] for i in range(len(firsts)-1))
check('citations ascending first-use order', ordered,
      'broken near ' + str(firsts[max(0, len(firsts)-6):]) if not ordered else '')
celltexts = [pp.text for t in doc.tables for r in t.rows for c in r.cells
             for pp in c.paragraphs if pp.text.strip()]
allcites = set(cites) | set(
    int(x) for c in re.findall(r'\[([\d,\s\u2013]+)\]', '\n'.join(celltexts))
    for x in re.split(r'[, \u2013]+', c) if x.isdigit())
check('all 1..43 cited (body or tables)', allcites == set(range(1, 44)),
      'missing: ' + str(sorted(set(range(1, 44)) - allcites)))

# ---- I. structure ----
titles = ['I. Introduction', 'II. Literature Review and Research Gap',
          'III. Device Structure and Material Selection', 'IV. Methodology',
          'V. Results and Discussion', 'VI. Comparison with Literature',
          'VII. Limitations and Future Work', 'VIII. Conclusion']
check('sections 1-8 titled', all(t in join_all for t in titles))
check('subsections A..T present', all(f'\n{n}. ' in '\n'+join_all or f' {n}. ' in join_all for n in 'ABCDEFGHIJKLMNOPQRST'))
check('Eqs. (1)-(3) referenced', 'Eqs. (1)\u2013(3)' in join_all)
for n in range(1, 21):
    check(f'Fig. {n} present', f'Fig. {n}.' in join_all or f'Fig. {n} ' in join_all)
nocap = [f'Table {n}' for n in ['I','II','III','IV','V','VI','VII','VIII'] if f'Table {n}.' not in join_all]
check('tables I-VIII captioned', not nocap, str(nocap))
check('no stale XIII/XIV refs', 'Table XIII' not in join_all and 'Table XIV' not in join_all)

# ---- J. terminology + numeric hygiene ----
check('no em dashes', '\u2014' not in text)
check('spike->step at RbGeI3/TiO2 (abstract+caption)', join_all.count('conduction-band step at RbGeI3/TiO2') == 2,
      str(join_all.count('conduction-band step at RbGeI3/TiO2')))
check('no spike at RbGeI3/TiO2 anywhere', 'spike at RbGeI3/TiO2' not in join_all)
check('26.80 appears 6+ times', text.count('26.80') >= 6, str(text.count('26.80')))
check('abstract no 19% claim', '19.' not in text[:1100])
check('NC = NV = 1x10^17 style', 'NC = NV = 1\u00d710\u00b9\u2077' in text)

# ---- K. format ----
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
pt = [t._tbl.find(f'{W}tblPr/{W}tblW') for t in doc.tables]
check(f'tables 100% width ({len(doc.tables)} tables)', len(pt) == 8 and all(tw is not None and tw.get(f'{W}w') == '5000' for tw in pt))
check('pages = 30', page_count() == 30, f'{page_count()} pages')
ps = [p for p in doc.paragraphs if p.text.strip().startswith('[')]
def ref_font(p):
    r = p._p.find(W+'r')
    rPr = r.find(W+'rPr') if r is not None else None
    rf = rPr.find(W+'rFonts') if rPr is not None else None
    return rf.get(W+'ascii') if rf is not None else None
fonts = {ref_font(p) for p in ps}
same = len(fonts) == 1 and None not in fonts
check('all refs share one font', same, f'fonts {fonts}')

print(f'\n===== {len(passes)} PASS, {len(fails)} FAIL =====')
for n, d in fails:
    print(f'FAIL  {n}  {d}')
sys.exit(1 if fails else 0)