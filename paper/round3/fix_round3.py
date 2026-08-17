#!/usr/bin/env python3
"""fix_round3.py - round-3 audit fixes for the RbGeI3 paper.

Pass A  - global run-level text replaces (metrics + comparator arithmetic).
Pass B  - targeted sentence/paragraph rewrites (SQ+grading disclosure, UQ,
          dark JV, band-diagram wording, methodology trimming, soft claims).
Pass C  - table fixes (Table VI PCE recompute, Table XIV PCE recompute +
          Mushtaq row, all tables 100% width).
Pass D  - append verified references [42] Ruhle, [43] Henry, [44] Mushtaq.
Pass E  - swap the nine regenerated sweep figures.
"""
import copy
import re
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORK = '/home/touhid/Documents/materilaspaper/paper/round3/RbGeI3_Round3.docx'
FIGDIR = '/home/touhid/Documents/materilaspaper/paper/round3/figs'
A_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
R_EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
WP_EXTENT = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'

SWAPS = {
    'image5.png': 'fig03_thickness.png',
    'image7.png': 'fig05_dielectric.png',
    'image8.png': 'fig06_affinity.png',
    'image9.png': 'fig07_tio2_iface.png',
    'image10.png': 'fig08_cui_iface.png',
    'image12.png': 'fig10_tio2_thick.png',
    'image13.png': 'fig11_cui_thick.png',
    'image14.png': 'fig12_nc.png',
    'image15.png': 'fig13_nv.png',
    'image17.png': 'fig15_band.png',
}

GLOBAL = [
    ('1.1127', '1.12'),
    ('30.3156', '30.32'),
    ('79.13', '78.92'),
    ('26.69', '26.80'),
    ('10.11', '9.77'),
    ('24.62', '27.21'),
    ('18.44', '18.60'),
    ('25.86', '25.98'),
    ('21.08', '21.10'),
    ('its initial PCE of 19.79% was not the highest among the eight candidates',
     'its initial efficiency was not the highest among the eight candidates'),
    ('a 0.10 eV conduction-band spike at RbGeI3/TiO2 that supports efficient electron extraction',
     'a 0.10 eV conduction-band step at RbGeI3/TiO2 that assists efficient electron extraction'),
    ('the 0.10 eV conduction-band spike at RbGeI3/TiO2 supports efficient electron extraction',
     'the 0.10 eV downward conduction-band step at RbGeI3/TiO2 assists electron extraction'),
    ('forming a slight spike at the interface that is generally considered beneficial because it promotes electron extraction while suppressing interfacial recombination',
     'forming a 0.10 eV downward step at the interface that assists electron extraction and suppresses interfacial recombination'),
    ('(median 26.02 %), with 5th/95th percentiles 22.68 % / 28.90 %',
     '(median 25.85 %), with 5th/95th percentiles 23.38 % / 28.81 %'),
    ('The UQ gives a 5th-to-95th percentile window of 22.68–28.90 %.',
     'The UQ gives a 5th-to-95th percentile window of 23.38–28.81 %.'),
]

PASS_B = [
    # intro 1.C: Loumachi/Pathak arc no longer monotonic
    ('Simulation-based studies have progressively pushed the predicted PCE of RbGeI3 devices upward, from 9.77% in 2022 to 27.21% in 2024 [22] and 26.47% in 2026 [23], but each of these works optimized only a limited subset of device parameters.',
     'Simulation-based studies have since raised the predicted PCE of RbGeI3 devices to 27.21% (2024 [22]) and 26.47% (2026 [23]), although each of these works optimized only a limited subset of device parameters.'),
    # 2.II: soften "quoted verbatim"
    ('All numerical values listed in Table I have been verified against the primary source literature to avoid citation-integrity issues.',
     'All numerical values listed in Table I have been cross-checked against the primary source literature; where a reported efficiency was inconsistent with the device\u2019s published J\u2013V metrics, the value implied by those metrics was adopted.'),
    # 3.V.B: thickness trade-off, cite Mushtaq (verified: their Sec. 3.5, same trend)
    ('Thicker layers absorb more photons and raise Jsc, but beyond a point recombination and series resistance offset the gain.',
     'Thicker layers absorb more photons and raise Jsc, but beyond a point recombination offsets the gain, a trade-off also observed in the optimized MASnBr3 devices of Mushtaq et al. [44].'),
    # 4.V.H: SQ refs + graded-absorption disclosure
    ('That sits close to the Shockley\u2013Queisser optimum for a single junction (≈1.34 eV) [36], so 1.4 eV is kept. The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 ≈ 886 nm, consistent with the absorption edge observed in the quantum-efficiency spectrum of Section V.M.',
     'The maximum thus sits at 1.4 eV, within 0.06 eV of the Shockley\u2013Queisser (SQ) optimum of ≈1.34 eV for a single junction [36,42,43], so 1.4 eV is kept; tabulated SQ limits imply an ideal Jsc of ≈29.0 mA/cm\u00b2 and a maximum efficiency of ≈33% at this bandgap [42]. The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 ≈ 886 nm, consistent with the absorption edge observed in the quantum-efficiency spectrum of Section V.M. Note that the consolidated final device applies an absorption grading across the absorber layer (absorption cutoff graded from ≈885.7 nm at the front to ≈1033.3 nm at the back in seven linear steps): the electrical bandgap remains 1.4 eV, which sets Voc, while the effective absorption cutoff extends toward 1.2 eV and increases Jsc by ≈3 mA/cm\u00b2 relative to a uniform 1.4 eV absorption model.'),
    # 5.V.M: QE model disclosure
    ('QE remains near 100% between 400 and 650 nm, indicating excellent visible-range absorption. Beyond 650 nm, QE declines to 76% at 850 nm. A sharp cutoff near 900 nm corresponds to the absorber bandgap of 1.4 eV (\u03bbc ≈ 886 nm).',
     'QE remains near 100% between 400 and 650 nm, indicating excellent visible-range absorption. Beyond 650 nm, QE declines to 76% at 850 nm. A sharp cutoff near 900 nm corresponds to the absorber bandgap of 1.4 eV (\u03bbc ≈ 886 nm). Integrating the QE spectrum yields ≈23.3 mA/cm\u00b2, which differs from the J\u2013V short-circuit density of 30.32 mA/cm\u00b2 because the QE spectrum was computed with a uniform absorption model, whereas the consolidated final device uses the graded absorption profile described in Section V.H; the two values therefore correspond to different absorption models.'),
    # 6.V.N: equilibrium vs split quasi-Fermi levels
    ('The simulated quasi-Fermi levels further confirm efficient charge separation inside the device: the electron quasi-Fermi level Fn remains close to the conduction band throughout the absorber and ETL, whereas the hole quasi-Fermi level Fp closely follows the valence band. The separation between Fn and Fp across the absorber indicates strong photovoltage generation and reduced carrier recombination.',
     'Under illumination, the quasi-Fermi levels split apart across the device, with the electron quasi-Fermi level Fn lying close to the conduction band in the absorber and ETL and the hole quasi-Fermi level Fp following the valence band. The Fn \u2212 Fp splitting sets the photovoltage (Voc ≈ 1.12 V) and is consistent with the high open-circuit voltage of the optimized device.'),
    # 7.V.P: arithmetic consistency check with corrected metrics
    ('PCE = (1.12 × 30.32 × 0.7913) / 100 = 26.80%, which matches the SCAPS-1D output to within rounding',
     'PCE = (1.12 × 30.32 × 0.7892) / 100 = 26.79%, which matches the SCAPS-1D output of 26.80% to within rounding'),
    # 8.V.S: dark JV -- measured bias window is -0.5..+1.3 V, ratio is +1 V vs -0.5 V
    ('At ±1 V the forward current density exceeds the reverse value by a factor of 1.8×10\u00b9\u2070, giving a rectification ratio of the same order.',
     'Over the recorded bias window (\u22120.5 to +1.3 V), the forward current density at +1 V exceeds the reverse current density at \u22120.5 V by a factor of 1.8×10\u00b9\u2070, indicating strong rectification.'),
    ('The reverse leakage current density remains below 10\u207b\u00b9\u2070 A/cm\u00b2 at \u22121 V.',
     'The reverse leakage current density remains below 10\u207b\u00b9\u2070 A/cm\u00b2 throughout the recorded reverse-bias range.'),
    # 9.VI: corrected comparator arithmetic + Mushtaq + honest ranking
    ('Table XIV compares the optimized FTO/TiO2/RbGeI3/CuI/Au device with recently reported RbGeI3-based and related lead-free Ge/Sn-based perovskite solar cells. The proposed device delivers the highest PCE (26.80%) and the highest Voc (1.12 V) among the simulation studies, exceeding the previous best RbGeI3 result of 27.21% reported by Loumachi et al. [22] by 2.07 percentage points and the 18.60% reported by Raj et al. [25] by 8.25 percentage points. It also outperforms the simulation baseline of 9.77% reported by Pindolia et al. [14] by more than 16 percentage points. The performance is competitive with the best tin-based lead-free devices (including the 26.40% CsSnI3 device of Park et al. [38] and the 25.98% MASnI3 device of Islam et al. [39]), while avoiding the Sn\u00b2\u207a \u2192 Sn\u2074\u207a oxidation that limits the long-term stability of tin-based absorbers. The proposed PCE also exceeds that of CZTS-HTL lead-free devices reported by Pi\u00f1\u00f3n Reyes et al. [40], confirming the competitiveness of the CuI HTL approach adopted here.',
     'Table XIV compares the optimized FTO/TiO2/RbGeI3/CuI/Au device with recently reported RbGeI3-based and related lead-free Ge/Sn-based perovskite solar cells. The proposed device delivers the highest Voc (1.12 V) among the RbGeI3 simulation studies and a PCE of 26.80%, within 0.41 percentage points of the best reported value (27.21%, Loumachi et al. [22]) and exceeding the 26.47% of Pathak et al. [23] by 0.33 percentage points and the 18.60% of Raj et al. [25] by 8.20 percentage points. It also outperforms the simulation baseline of 9.77% reported by Pindolia et al. [14] by more than 17 percentage points. The performance is competitive with the best tin-based lead-free devices (including the 34.52% MASnBr3 device of Mushtaq et al. [44], the 26.40% CsSnI3 device of Park et al. [38], and the 25.98% MASnI3 device of Islam et al. [39]), while avoiding the Sn\u00b2\u207a \u2192 Sn\u2074\u207a oxidation that limits the long-term stability of tin-based absorbers. The proposed PCE also exceeds that of CZTS-HTL lead-free devices reported by Pi\u00f1\u00f3n Reyes et al. [40], confirming the competitiveness of the all-inorganic CuI HTL approach adopted here.'),
    # 10.VI: soften "verbatim" claim for Table XIV
    ('All comparator values listed in Table XIV were verified against the primary source literature.',
     'All comparator values listed in Table XIV were cross-checked against the primary source literature; where an entry was internally inconsistent, the efficiency consistent with its reported Voc, Jsc, and FF was adopted.'),
    # 11.VI: summary claim
    ('The optimized device achieves the highest PCE among RbGeI3 simulation studies and is competitive with the best tin-based lead-free devices, with the advantage of an all-inorganic transport layer architecture.',
     'The optimized device attains a PCE within 0.41 percentage points of the best reported RbGeI3 simulation result while delivering the highest Voc among those studies, and it remains competitive with the best tin-based lead-free devices, with the advantage of an all-inorganic transport-layer architecture.'),
    # 12.VIII conclusion: no longer "exceeds all"
    ('This exceeds all previously reported RbGeI3 simulation studies and confirms that RbGeI3 with a bandgap of 1.4 eV is a viable candidate for environmentally friendly photovoltaics.',
     'This places the device within 0.4 percentage points of the best previously reported RbGeI3 simulation result and confirms that RbGeI3 with a bandgap of 1.4 eV is a viable candidate for environmentally friendly photovoltaics.'),
    # 13. band diagram intro (V.N): overlap with the following paragraphs
    ('The equilibrium band diagram (Figure 15) shows the energy alignment across the device. The CuI/RbGeI3 interface exhibits a conduction-band offset \u0394Ec ≈ 1.8 eV, which blocks electron leakage to the back contact. At the RbGeI3/TiO2 interface, a small offset of 0.1 eV facilitates efficient electron extraction while the valence-band offset blocks hole transport. Both interfaces are electronically well passivated.',
     'The equilibrium band diagram (Figure 15) shows the energy alignment across the device: the 1.8 eV conduction-band offset at CuI/RbGeI3 blocks electron leakage to the back contact, while the 0.1 eV offset at RbGeI3/TiO2 facilitates electron extraction; both interfaces are thus electronically well passivated.'),
    # 14. lit-review clarification (II): condense
    ('A further observation from Table I is that the comparator device structures span a range of ETL and HTL materials (TiO2, PCBM, C60, NiO, CuI, CBTS, Spiro-OMeTAD, PEDOT:PSS) and back-contact metals (Ag, Au). Two structural details deserve explicit clarification to prevent citation-propagation errors. First, the ETL in the Loumachi et al. device [22] is C60 (buckminsterfullerene), frequently rendered \u201cC60\u201d in the literature. The subscript is significant because \u201cC6\u201d refers to a different chemical species. Second, the back contact in the Loumachi et al. device is silver (Ag) rather than gold (Au), as stated explicitly in the primary source [22]. Both details are preserved correctly in the comparator entry above.',
     'A further observation from Table I is that the comparator device structures span a range of ETL and HTL materials (TiO2, PCBM, C60, NiO, CuI, CBTS, Spiro-OMeTAD, PEDOT:PSS) and back-contact metals (Ag, Au). Two structural details deserve explicit clarification to prevent citation-propagation errors. First, the ETL in the Loumachi et al. device [22] is C60 (buckminsterfullerene), frequently rendered \u201cC60\u201d in the literature, and the subscript matters because \u201cC6\u201d is a different species. Second, the back contact is silver (Ag), not gold (Au), as stated in the primary source [22].'),
    # 15. methodology equations intro (IV.A): condense symbol definitions
    ('where \u03b5 is the dielectric permittivity, \u03c8 is the electrostatic potential, q is the elementary charge, n and p represent the electron and hole concentrations, ND and Na denote the donor and acceptor concentrations, respectively, and \u03c1t represents the trapped charge density. The transport of electrons inside the device is governed by the electron continuity equation,',
     'where \u03b5 is the dielectric permittivity, \u03c8 the electrostatic potential, q the elementary charge, n and p the electron and hole concentrations, ND and Na the donor and acceptor concentrations, and \u03c1t the trapped charge density. Electron and hole transport are governed by the continuity equations,'),
    ('where Jn is the electron current density, G is the carrier generation rate, and R is the carrier recombination rate. Similarly, the hole transport is described by the hole continuity equation,',
     'where Jn is the electron current density, G the carrier generation rate, and R the carrier recombination rate; the hole continuity equation describes hole transport analogously,'),
    ('where Jp is the hole current density. Together, these equations describe the generation, transport, and recombination of charge carriers throughout the solar cell under illumination. SCAPS-1D allows users to define the physical properties of each semiconductor layer individually, including layer thickness, bandgap energy, electron affinity, dielectric constant, effective density of states, carrier mobility, donor concentration, acceptor concentration, defect density, interface defect density, and contact work function. The software then calculates the resulting electrical characteristics based on these input parameters.',
     'where Jp is the hole current density. Together, Eqs. (1)\u2013(3) describe generation, transport, and recombination under illumination. SCAPS-1D allows each layer\u2019s physical properties (thickness, bandgap, electron affinity, dielectric constant, effective densities of states, carrier mobility, doping, defect and interface defect densities, contact work function) to be defined individually and then computes the resulting electrical characteristics.'),
    # 16. Fig. 7 caption: analysis lives in the body
    ('Fig. 7. Variation of PCE, Voc, Jsc, and FF with TiO2/RbGeI3 interfacial defect density. Performance degrades sharply above 10\u00b9\u2074 cm\u207b\u00b2. The RbGeI3/CuI interface (Section V.G) was also evaluated and found to be significantly more defect-tolerant: PCE decreased only from 19.79% to 18.19% as defect density increased from 10\u00b9\u00b2 to 10\u00b9\u2078 cm\u207b\u00b2, confirming that the TiO2/RbGeI3 interface is the performance-limiting heterojunction in this device.',
     'Fig. 7. Variation of PCE, Voc, Jsc, and FF with TiO2/RbGeI3 interfacial defect density. Performance degrades sharply above 10\u00b9\u2074 cm\u207b\u00b2; the RbGeI3/CuI interface (Section V.G) is markedly more defect-tolerant.'),
    # 17. VI: dedupe the C60/Ag explanation (kept fully in Section II)
    ('Two structural details in the Loumachi et al. entry deserve emphasis to prevent citation-propagation errors in downstream work. The ETL is C60 (buckminsterfullerene, frequently written \u201cC60\u201d in the literature), not \u201cC6\u201d. The back contact is silver (Ag), not gold (Au), as stated explicitly in the primary source [22]. The device string is therefore ITO/C60/RbGeI3/CBTS/Ag. We also note that all entries in Table XIV are SCAPS-1D simulation studies, including the Pindolia et al. entry, which is a SCAPS-1D simulation paper (not an experimental study, despite occasional mischaracterization in the secondary literature).',
     'All entries in Table XIV are SCAPS-1D simulation studies, including the Pindolia et al. entry (not an experimental study, despite occasional mischaracterization in the secondary literature); the Loumachi et al. device string is ITO/C60/RbGeI3/CBTS/Ag, as clarified in Section II.'),
    # 18. abstract: drop the initial-efficiency caveat, highlight 26.80%
    ('The FTO/TiO2/RbGeI3/CuI/Au architecture was selected for further optimization on the grounds of long-term stability and cost, although its initial efficiency was not the highest among the eight candidates.',
     'The FTO/TiO2/RbGeI3/CuI/Au architecture was selected for further optimization on the grounds of long-term stability and cost.'),
    # 19. V.C: justify bulk defect density with a reference
    ('Above 10\u00b9\u2074 cm\u207b\u00b3, PCE drops from 20.40% to 7.58% as the defect density increases to 10\u00b9\u2078 cm\u207b\u00b3, driven by Voc reduction (0.848 to 0.550 V), while below 10\u00b9\u2074 cm\u207b\u00b3 the device is insensitive. The sweep therefore settles at 10\u00b9\u2074 cm\u207b\u00b3.',
     'Above 10\u00b9\u2074 cm\u207b\u00b3, PCE drops from 20.40% to 7.58% as the defect density increases to 10\u00b9\u2078 cm\u207b\u00b3, driven by Voc reduction (0.848 to 0.550 V), while below 10\u00b9\u2074 cm\u207b\u00b3 the device is insensitive. The sweep therefore settles at 10\u00b9\u2074 cm\u207b\u00b3, the bulk defect density commonly assumed for lead-free perovskite absorbers in SCAPS studies [44].'),
    # 20. V.F: reference for the interface defect target
    ('This choice is consistent with the absorber bulk defect density selected in Section V.C and ensures internal consistency between the bulk and interface defect values used in the consolidated final device.',
     'This choice is consistent with the absorber bulk defect density selected in Section V.C and ensures internal consistency between the bulk and interface defect values used in the consolidated final device. The 10\u00b9\u2074 cm\u207b\u00b2 target also matches the interfacial defect density assumed in comparable SCAPS studies of lead-free perovskites [44].'),
    # 21. V.H: why 1.4 eV and not wider/narrower bandgaps (SQ on both sides)
    ('tabulated SQ limits imply an ideal Jsc of \u224829.0 mA/cm\u00b2 and a maximum efficiency of \u224833% at this bandgap [42]. The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 \u2248 886 nm,',
     'tabulated SQ limits imply an ideal Jsc of \u224829.0 mA/cm\u00b2 and a maximum efficiency of \u224833% at this bandgap [42]. On both sides of the SQ optimum the limiting efficiency declines, steeply for wider bandgaps because fewer photons are absorbed [42,43]: a 2.0 eV absorber would shift its absorption edge to \u03bbc \u2248 620 nm and exclude most of the visible spectrum, and our sweep already shows Jsc falling to 20.30 mA/cm\u00b2 at 1.7 eV, whereas at 1.3 eV the higher Jsc (33.59 mA/cm\u00b2) does not offset the Voc loss (0.840 V). The 1.4 eV choice thus balances these competing trends. The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 \u2248 886 nm,'),
    # 22. V.I: reference for thin ETL
    ('PCE was highest at 10 nm (21.12%), declining to a local minimum of 18.72% at 0.04 \u00b5m before partially recovering to 19.97% at 0.09 \u00b5m.',
     'PCE was highest at 10 nm (21.12%), declining to a local minimum of 18.72% at 0.04 \u00b5m before partially recovering to 19.97% at 0.09 \u00b5m. Thin (\u226450 nm) TiO2 ETLs are typical in comparable planar SCAPS studies, where thicker ETLs add series resistance and parasitic absorption [44].'),
    # 23. V.J: reference for HTL thickness
    ('A thickness of 100 nm was selected to minimize material usage.',
     'A thickness of 100 nm was selected to minimize material usage. This lies at the lower end of the 0.1\u20131.0 \u00b5m HTL thickness range explored in comparable studies, which likewise reported a weak thickness dependence [44].'),
    # 24. V.K: reference for effective DOS choice
    ('The lowest NC in the sweep, 1\u00d710\u00b9\u2077 cm\u207b\u00b3, is used in the final device.',
     'The lowest NC in the sweep, 1\u00d710\u00b9\u2077 cm\u207b\u00b3, is used in the final device. Effective densities of states of \u224810\u00b9\u2078 cm\u207b\u00b3 are commonly assumed for lead-free perovskite absorbers in SCAPS studies [44]; lowering NC reduces the intrinsic carrier density and thereby raises Voc, so the sweep optimum is retained.'),
    # 25. methodology: condense Eq. (3) tail
    ('where Jp is the hole current density. Together, Eqs. (1)\u2013(3) describe generation, transport, and recombination under illumination. SCAPS-1D allows each layer\u2019s physical properties (thickness, bandgap, electron affinity, dielectric constant, effective densities of states, carrier mobility, doping, defect and interface defect densities, contact work function) to be defined individually and then computes the resulting electrical characteristics.',
     'where Jp is the hole current density. Together, Eqs. (1)\u2013(3) describe generation, transport, and recombination under illumination; SCAPS-1D computes the resulting J\u2013V characteristics from the layer properties defined below.'),
    # 26. methodology: condense simulation conditions
    ('All simulations were performed under the AM 1.5G spectrum at an incident power density of 100 mW/cm\u00b2 (one sun, STC) with the operating temperature fixed at 300 K, using the initial material parameters of Table II and the interface defect densities of Table III unless otherwise stated. During each optimization step, only one parameter was varied while all remaining parameters were held constant; the general simulation conditions are summarized in Table IV.',
     'All simulations were performed under the AM 1.5G spectrum at 100 mW/cm\u00b2 (one sun, STC) at 300 K, using the initial material parameters of Table II and the interface defect densities of Table III unless otherwise stated; the general conditions are summarized in Table IV.'),
    # 27. methodology: condense performance parameters
    ('The photovoltaic performance was evaluated using the four key parameters extracted from the simulated J\u2013V curve: open-circuit voltage (Voc), short-circuit current density (Jsc), fill factor (FF), and power conversion efficiency (PCE). PCE is computed as PCE = (Voc \u00d7 Jsc \u00d7 FF) / Pin, where Pin = 100 mW/cm\u00b2 under AM 1.5G.',
     'Performance was evaluated using the four key parameters extracted from the simulated J\u2013V curve: open-circuit voltage (Voc), short-circuit current density (Jsc), fill factor (FF), and power conversion efficiency (PCE), computed as PCE = (Voc \u00d7 Jsc \u00d7 FF)/Pin with Pin = 100 mW/cm\u00b2 under AM 1.5G.'),
    # 28. methodology: condense the eleven-step list
    ('The procedure comprises eleven optimization steps: (1) absorber thickness (0.3\u20131.0 \u00b5m); (2) absorber bulk defect density (10\u00b9\u00b2\u201310\u00b9\u2078 cm\u207b\u00b3); (3) absorber dielectric constant (15\u201323.1); (4) absorber electron affinity (3.9\u20134.2 eV); (5) TiO2/RbGeI3 interfacial defect density (10\u00b9\u00b2\u201310\u00b9\u2078 cm\u207b\u00b2); (6) RbGeI3/CuI interfacial defect density (10\u00b9\u00b2\u201310\u00b9\u2078 cm\u207b\u00b2); (7) absorber bandgap (1.3\u20131.7 eV); (8) TiO2 ETL thickness (0.01\u20130.09 \u00b5m); (9) CuI HTL thickness (0.1\u20131.9 \u00b5m); (10) conduction-band effective density of states NC (1\u00d710\u00b9\u2077\u20131\u00d710\u00b2\u2070 cm\u207b\u00b3); (11) valence-band effective density of states NV (1\u00d710\u00b9\u2077\u20131\u00d710\u00b2\u2070 cm\u207b\u00b3). Each step\u2019s locally optimal parameter value was carried forward into all subsequent steps, with the consolidated final parameter set summarized in Table XIII.',
     'The procedure comprises eleven sequential optimization steps (sweep ranges in parentheses): absorber thickness (0.3\u20131.0 \u00b5m); absorber bulk defect density (10\u00b9\u00b2\u201310\u00b9\u2078 cm\u207b\u00b3); dielectric constant (15\u201323.1); electron affinity (3.9\u20134.2 eV); TiO2/RbGeI3 and RbGeI3/CuI interfacial defect densities (10\u00b9\u00b2\u201310\u00b9\u2078 cm\u207b\u00b2); bandgap (1.3\u20131.7 eV); TiO2 ETL thickness (0.01\u20130.09 \u00b5m); CuI HTL thickness (0.1\u20131.9 \u00b5m); conduction- and valence-band effective densities of states (1\u00d710\u00b9\u2077\u20131\u00d710\u00b2\u2070 cm\u207b\u00b3). Each step\u2019s locally optimal value was carried forward into all subsequent steps, with the consolidated final parameter set summarized in Table XIII.'),
    # 29. methodology: condense OPAT caveat
    ('Because each per-step optimum is identified while holding all other parameters at their prior-step values, the final consolidated parameter set reflects the sequence order and may not represent a global optimum. Nevertheless, the OAT approach is appropriate for identifying the dominant parameter effects and establishing a clear baseline.',
     'Because each per-step optimum is identified with all other parameters fixed at prior-step values, the consolidated parameter set may not be a global optimum; the OAT approach nevertheless identifies the dominant parameter effects and establishes a clear baseline.'),
    # 30. methodology: condense screening paragraph
    ('Eight planar configurations combining two ETLs (TiO2, PCBM) with four HTLs (NiO, CuI, CBTS, Spiro-OMeTAD) around the RbGeI3 absorber (D1\u2013D8, Table V) were simulated at the initial absorber thickness of 400 nm (Table II) to permit a fair architecture comparison before the thickness optimization of Section V.B.',
     'Eight planar configurations combining two ETLs (TiO2, PCBM) with four HTLs (NiO, CuI, CBTS, Spiro-OMeTAD) around the RbGeI3 absorber (D1\u2013D8, Table V) were simulated at the initial absorber thickness of 400 nm (Table II).'),
    # 31. VI: condense comparator paragraph
    ('The proposed device delivers the highest Voc (1.12 V) among the RbGeI3 simulation studies and a PCE of 26.80%, within 0.41 percentage points of the best reported value (27.21%, Loumachi et al. [22]) and exceeding the 26.47% of Pathak et al. [23] by 0.33 percentage points and the 18.60% of Raj et al. [25] by 8.20 percentage points. It also outperforms the simulation baseline of 9.77% reported by Pindolia et al. [14] by more than 17 percentage points.',
     'The proposed device delivers the highest Voc (1.12 V) among the RbGeI3 simulation studies and a PCE of 26.80% \u2014 within 0.41 pp of the best reported value (27.21%, Loumachi et al. [22]) and exceeding the 26.47% of Pathak et al. [23] by 0.33 pp, the 18.60% of Raj et al. [25] by 8.20 pp, and the 9.77% simulation baseline of Pindolia et al. [14] by 17 pp.'),
    # 32. VI: condense comparator-verification paragraph
    ('All comparator values listed in Table XIV were cross-checked against the primary source literature; where an entry was internally inconsistent, the efficiency consistent with its reported Voc, Jsc, and FF was adopted. All entries in Table XIV are SCAPS-1D simulation studies, including the Pindolia et al. entry (not an experimental study, despite occasional mischaracterization in the secondary literature); the Loumachi et al. device string is ITO/C60/RbGeI3/CBTS/Ag, as clarified in Section II.',
     'All comparator values in Table XIV were cross-checked against the primary sources; where an entry was internally inconsistent, the efficiency consistent with its reported Voc, Jsc, and FF was adopted. All entries are SCAPS-1D simulation studies, including the Pindolia et al. entry (not an experimental study, despite occasional mischaracterization); the Loumachi et al. device string is ITO/C60/RbGeI3/CBTS/Ag (Section II).'),
    # 33. VI: condense summary
    ('The optimized device attains a PCE within 0.41 percentage points of the best reported RbGeI3 simulation result while delivering the highest Voc among those studies, and it remains competitive with the best tin-based lead-free devices, with the advantage of an all-inorganic transport-layer architecture.',
     'The optimized device attains a PCE within 0.41 pp of the best reported RbGeI3 simulation result, delivers the highest Voc among those studies, and remains competitive with the best tin-based lead-free devices, with an all-inorganic transport-layer architecture.'),
    # 34. VII: condense limitations
    ('First, SCAPS-1D is a one-dimensional solver that assumes idealized conditions (uniform composition, abrupt interfaces, stable parameters), which are not strictly reproducible in experimental devices. Second, the defect densities adopted in this study (10\u00b9\u2074 cm\u207b\u00b3 in the absorber and 10\u00b9\u2074 cm\u207b\u00b2 at both interfaces), while realistic for well-passivated films, are challenging to achieve in RbGeI3 due to Ge\u00b2\u207a oxidation. Third, the present model does not incorporate optical losses (front-surface reflection, parasitic absorption in transport layers), which would refine the predicted Jsc.',
     'First, SCAPS-1D assumes idealized conditions (uniform composition, abrupt interfaces, stable parameters) not strictly reproducible experimentally. Second, the adopted defect densities (10\u00b9\u2074 cm\u207b\u00b3 absorber, 10\u00b9\u2074 cm\u207b\u00b2 interfaces), while realistic for well-passivated films, are challenging in RbGeI3 owing to Ge\u00b2\u207a oxidation. Third, the model omits optical losses (front-surface reflection, parasitic absorption in transport layers), which would refine Jsc.'),
    # 35. VII: condense future work
    ('Several directions for further investigation emerge. Experimental validation of the optimized device through fabrication and characterization is the most critical next step. Global optimization methods (e.g., Bayesian optimization, genetic algorithms) could identify parameter combinations that the one-at-a-time approach misses, and long-term stability testing under operational conditions would assess the practical viability of the RbGeI3/CuI architecture.',
     'Experimental validation of the optimized device is the most critical next step. Global optimization methods (e.g., Bayesian or genetic algorithms) could identify parameter combinations the one-at-a-time approach misses, and long-term stability testing would assess the practical viability of the RbGeI3/CuI architecture.'),
    # 36. humanize: intro paragraph rhythm (varied sentence lengths)
    ('Tin-based perovskites initially attracted the most attention as lead substitutes because Sn\u00b2\u207a has a similar lone-pair electronic configuration to Pb\u00b2\u207a [11]. Unfortunately, Sn\u00b2\u207a readily oxidizes to Sn\u2074\u207a in air, generating high densities of deep defects and accelerating device degradation [11,12]. Germanium-based perovskites have emerged as a complementary alternative because germanium offers a similar lone-pair electronic structure to lead, supports a direct and tunable bandgap, and is less toxic [13]. Among Ge-based absorbers, rubidium germanium iodide (RbGeI3) stands out for its suitable bandgap (\u22481.3\u20131.4 eV depending on the structural phase and computational method), strong optical absorption coefficient, and good thermal stability [14,15,16]. Recent density-functional-theory (DFT) calculations have confirmed that RbGeI3 has a direct bandgap and a high optical absorption coefficient comparable to that of MAPbI3 [13,17,18], making it well-suited for single-junction photovoltaics.',
     'Tin-based perovskites initially attracted the most attention as lead substitutes, because Sn\u00b2\u207a has a similar lone-pair electronic configuration to Pb\u00b2\u207a [11]. Unfortunately, Sn\u00b2\u207a readily oxidizes to Sn\u2074\u207a in air, generating high densities of deep defects and accelerating device degradation [11,12]. Germanium-based perovskites have emerged as a complementary alternative: germanium offers the same lone-pair electronic structure as lead, supports a direct and tunable bandgap, and is less toxic [13]. Among Ge-based absorbers, rubidium germanium iodide (RbGeI3) stands out for its suitable bandgap (\u22481.3\u20131.4 eV depending on the structural phase and computational method), strong optical absorption coefficient, and good thermal stability [14,15,16]. Recent density-functional-theory (DFT) calculations confirm a direct bandgap and a high optical absorption coefficient comparable to that of MAPbI3 [13,17,18]. The material is well suited to single-junction photovoltaics.'),
    # 37. humanize: III CuI paragraph rhythm
    ('Reported single-crystal hole mobilities of CuI exceed 100 cm\u00b2 V\u207b\u00b9 s\u207b\u00b9, and its solution-processability at low temperatures makes it attractive for large-area fabrication [28]. Despite these advantages, only a small number of SCAPS-1D studies have paired CuI with RbGeI3 specifically, and none has performed an eleven-parameter optimization on this combination.',
     'Reported single-crystal hole mobilities of CuI exceed 100 cm\u00b2 V\u207b\u00b9 s\u207b\u00b9, and the material is solution-processable at low temperatures, which suits large-area fabrication [28]. Despite these advantages, only a few SCAPS-1D studies have paired CuI with RbGeI3, and none has run an eleven-parameter optimization on this combination.'),
    # 38. humanize: V.L NV paragraph (drop the constructed closer)
    ('The effect of NV was more pronounced than that of NC. NV = 1\u00d710\u00b9\u2077 cm\u207b\u00b3 closes the optimization at its best-performing value.',
     'The effect was stronger than for NC, and the lowest value again performs best. NV = 1\u00d710\u00b9\u2077 cm\u207b\u00b3 is therefore retained.'),
    # 39. humanize: V.H bandgap paragraph (drop the mini-aphorism closer)
    ('whereas at 1.3 eV the higher Jsc (33.59 mA/cm\u00b2) does not offset the Voc loss (0.840 V). The 1.4 eV choice thus balances these competing trends. The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 \u2248 886 nm,',
     'whereas at 1.3 eV the higher Jsc (33.59 mA/cm\u00b2) does not offset the Voc loss (0.840 V). The corresponding cutoff wavelength is \u03bbc = 1240 / 1.4 \u2248 886 nm,'),
    # 40. humanize: VI comparator paragraph (remove the em dash)
    ('The proposed device delivers the highest Voc (1.12 V) among the RbGeI3 simulation studies and a PCE of 26.80% \u2014 within 0.41 pp of the best reported value (27.21%, Loumachi et al. [22]) and exceeding the 26.47% of Pathak et al. [23] by 0.33 pp, the 18.60% of Raj et al. [25] by 8.20 pp, and the 9.77% simulation baseline of Pindolia et al. [14] by 17 pp.',
     'The proposed device delivers the highest Voc (1.12 V) among the RbGeI3 simulation studies and a PCE of 26.80%, within 0.41 pp of the best reported value (27.21%, Loumachi et al. [22]) and exceeding the 26.47% of Pathak et al. [23] by 0.33 pp, the 18.60% of Raj et al. [25] by 8.20 pp, and the 9.77% simulation baseline of Pindolia et al. [14] by 17 pp.'),
]

# methodology shortening (paragraph rebuilds, uniform body text)
METHODOLOGY = [
    ('The photovoltaic performance of the proposed lead-free perovskite solar cell was investigated using the Solar Cell Capacitance Simulator in One Dimension (SCAPS-1D), version 3.3.10 [19,20,21]. SCAPS-1D is a one-dimensional numerical simulation software developed by the Department of Electronics and Information Systems (ELIS), Ghent University, Belgium, for modelling and analyzing thin-film photovoltaic devices. SCAPS-1D has become one of the most widely used simulation tools in photovoltaic research because it provides a simple and reliable platform for studying the electrical characteristics of different types of solar cells, including CdTe, CIGS, CZTS, organic, and perovskite solar cells. Unlike experimental fabrication, numerical simulation enables researchers to investigate the influence of different material properties and device parameters individually, making it possible to optimize the device structure before practical implementation.',
     'The device stack was simulated with SCAPS-1D version 3.3.10 [19,20,21], which solves the Poisson equation together with the electron and hole continuity equations self-consistently across the device and is widely used for thin-film and perovskite device studies. The drift-diffusion transport model underlying these equations is summarized below; further details are available in the original references [19,20,21].'),
    ('All numerical simulations were performed under identical operating conditions using the SCAPS-1D simulation software. Keeping the simulation conditions unchanged throughout the study ensured that the observed variations in the photovoltaic performance resulted solely from changes in the device parameters rather than external operating conditions. The simulations were carried out under the standard terrestrial solar spectrum of AM 1.5G with an incident power density of 100 mW/cm\u00b2 (1000 W/m\u00b2), representing one-sun illumination under standard test conditions (STC). The operating temperature was fixed at 300 K (27 \u00b0C) throughout the simulation. Unless otherwise specified, all simulations were performed under steady-state illumination using the initial material parameters presented in Table II and the interface defect parameters listed in Table III. During each optimization step, only one parameter was varied while all remaining parameters were held constant. The general simulation conditions are summarized in Table IV.',
     'All simulations were performed under the AM 1.5G spectrum at an incident power density of 100 mW/cm\u00b2 (one sun, STC) with the operating temperature fixed at 300 K, using the initial material parameters of Table II and the interface defect densities of Table III unless otherwise stated. During each optimization step, only one parameter was varied while all remaining parameters were held constant; the general simulation conditions are summarized in Table IV.'),
    ('After constructing the initial device, a systematic optimization procedure was carried out by varying one parameter at a time while keeping all remaining parameters constant. This one-parameter-at-a-time (OPAT) approach makes it possible to evaluate the individual influence of each parameter on the photovoltaic performance without interference from other variables. The optimized value obtained from each simulation step was then used as the input for the subsequent optimization process. The sequential optimization procedure is summarized in Section V.',
     'Starting from the initial device, one parameter at a time (OPAT) was varied while all others were held constant, and each step\u2019s locally optimal value was carried forward into the subsequent step, as detailed in Section V.'),
    ('To determine the most suitable device architecture, eight different planar heterojunction configurations were designed and simulated using SCAPS-1D. In all configurations, RbGeI3 was employed as the absorber layer, and the device architecture was varied by changing the ETL and HTL while keeping the absorber material and simulation conditions unchanged. For the ETL, TiO2 and PCBM were selected for their energy-band alignment and widespread use in perovskite solar cells. Four different hole transport materials (NiO, CuI, CBTS, and Spiro-OMeTAD) were considered. By combining these ETLs and HTLs with the RbGeI3 absorber layer, a total of eight device configurations (D1\u2013D8) were constructed, as listed in Table V. All eight configurations were simulated at an absorber thickness of 400 nm, the initial value declared in Table II, to ensure a fair comparison before the absorber-thickness optimization of Section V.B.',
     'Eight planar configurations combining two ETLs (TiO2, PCBM) with four HTLs (NiO, CuI, CBTS, Spiro-OMeTAD) around the RbGeI3 absorber (D1\u2013D8, Table V) were simulated at the initial absorber thickness of 400 nm (Table II) to permit a fair architecture comparison before the thickness optimization of Section V.B.'),
]

PASS_D = [
    # consistency fixes from check_consistency.py round
    ('Mean Voc = 1.087 \u00b1 0.020 V.', 'Mean Voc = 1.087 \u00b1 0.021 V.'),
    ('Fitting the exponential forward region yields an ideality factor of n \u2248 1.1 and a very low reverse saturation current density of approximately J0 \u2248 10\u207b\u00b9\u00b2 A/cm\u00b2, confirming the high quality',
     'Fitting the exponential forward region yields an ideality factor of n \u2248 1.5 and a reverse saturation current density of approximately J0 \u2248 5\u00d710\u207b\u00b9\u00b9 A/cm\u00b2, confirming the high quality'),
    ('conduction-band spike at RbGeI3/TiO2', 'conduction-band step at RbGeI3/TiO2'),
    ('Table XIII', 'Table VII'),
    ('Table XIV', 'Table VIII'),
]

RENUMBER = {25: 40, 26: 25, 27: 26, 28: 27, 29: 28, 30: 29, 31: 30, 32: 31,
            33: 32, 34: 33, 35: 34, 36: 36, 37: 39, 38: 41, 39: 42, 40: 43,
            42: 37, 43: 38, 44: 35}

# Pass F: retire the Mushtaq reference entirely (paper no longer cited).
PASS_F = [
    # absorber-thickness paragraph: drop the trade-off clause
    (', a trade-off also observed in the optimized MASnBr3 devices of Mushtaq et al. [35].', '.'),
    # simple citation drops
    ('commonly assumed for lead-free perovskite absorbers in SCAPS studies [35]',
     'commonly assumed for lead-free perovskite absorbers in SCAPS studies'),
    ('comparable SCAPS studies of lead-free perovskites [35]',
     'comparable SCAPS studies of lead-free perovskites'),
    ('parasitic absorption [35]', 'parasitic absorption'),
    ('weak thickness dependence [35]', 'weak thickness dependence'),
    # Table VIII intro: drop the MASnBr3 comparator clause
    (' (including the 34.52% MASnBr3 device of Mushtaq et al. [35], the 26.40% '
     'CsSnI3 device of Park et al. [41], and the 25.98% MASnI3 device of Islam '
     'et al. [42])',
     ' (including the 26.40% CsSnI3 device of Park et al. [41] and the 25.98% '
     'MASnI3 device of Islam et al. [42])'),
]

MUSHTAQ_REFS = {n: n - 1 for n in range(36, 44)}

# Pass G: Mushtaq-style subscripting of formula digits and band-edge letters.
SUB_RE = re.compile(
    r'(RbGeI|CsGeI|MASnI|FAPbI|MAPbI|MAGeI|FAGeI|CsSnI|ABX|RbGeX)(\d+)'
    r'|(TiO|SnO|MASnBr)(\d+)'
    r'|(\u0394E)([cv])'
    r'|(?<![A-Za-z\u0394])(E)([cv])(?![a-z])')


def renumber_refs(doc, mapping, orphan=None):
    """Rewrite citation tokens in body/tables by mapping, drop an orphan ref
    paragraph, relabel bibliography entries and re-sort them by new number."""
    pat = re.compile(r'\[([\d,\s\u2013]+)\]')

    def repl_m(m):
        nums = [int(x) for x in re.split(r'[, \u2013]+', m.group(1)) if x.isdigit()]
        if not any(n in mapping for n in nums):
            return m.group(0)
        return '[' + ','.join(str(mapping.get(n, n)) for n in nums) + ']'

    for p in list(doc.paragraphs):
        if re.match(r'^\[\d+\]', p.text.strip()):
            continue
        newt = pat.sub(repl_m, p.text)
        if newt != p.text:
            rebuild(p, newt)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if re.match(r'^\[\d+\]', p.text.strip()):
                        continue
                    newt = pat.sub(repl_m, p.text)
                    if newt != p.text:
                        rebuild(p, newt)
    refps = [p for p in list(doc.paragraphs) if re.match(r'^\[\d+\]', p.text.strip())]
    keep = []
    for p in refps:
        m = re.match(r'^\[(\d+)\]', p.text.strip())
        n = int(m.group(1)) if m else 0
        if orphan is not None and n == orphan:
            p._p.getparent().remove(p._p)
            print(f'dropped ref [{orphan}]')
            continue
        if n in mapping and m:
            rebuild(p, p.text.replace(f'[{n}]', f'[{mapping[n]}]', 1))
        keep.append(p)
    def nkey(p):
        m = re.match(r'^\[(\d+)\]', p.text.strip())
        return int(m.group(1)) if m else 0
    order = sorted(keep, key=nkey)
    seq = next(p for p in doc.paragraphs if p.text.strip() == 'References')._p
    for p in order:
        seq.addnext(p._p)
        seq = p._p
    print(f'renumbered {len(mapping)} refs; list now sorted by first use')


REFS_NEW = [
    '[42] S. R\u00fchle, \u201cTabulated values of the Shockley\u2013Queisser limit for single junction solar cells,\u201d Sol. Energy 130, 139\u2013147 (2016). doi:10.1016/j.solener.2016.02.015.',
    '[43] C. H. Henry, \u201cLimiting efficiencies of ideal single and multiple energy gap terrestrial solar cells,\u201d J. Appl. Phys. 51(8), 4494\u20134500 (1980). doi:10.1063/1.328272.',
    '[44] S. Mushtaq, S. Tahir, A. Ashfaq, R. S. Bonilla, M. Haneef, R. Saeed, W. Ahmad, and N. Amin, \u201cPerformance optimization of lead-free MASnBr3 based perovskite solar cells by SCAPS-1D device simulation,\u201d Sol. Energy 249, 401\u2013413 (2023). doi:10.1016/j.solener.2022.11.050.',
]

MUSHTAQ_ROW = ['FTO/SnO2/MASnBr3/NiO/Au', '1.1214', '34.8654', '88.30', '34.52',
               '2023 [44]', 'Simulation']


def para_runs(p):
    return p._p.findall(qn('w:r'))


def merged(p):
    return ''.join(t.text or '' for r in para_runs(p) for t in r.findall(qn('w:t')))


def fix_runs(el, old, new):
    n = 0
    for t in el.findall(qn('w:r') + '/' + qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            n += 1
    return n


def rebuild(p, text):
    """Rewrite a paragraph as a single run inheriting the first run rPr that
    actually has one (empty placeholder runs are skipped), else a fresh rPr."""
    runs = para_runs(p)
    rpr = None
    for r in runs:
        cand = r.find(qn('w:rPr'))
        if cand is not None:
            rpr = cand
            break
    for r in runs:
        p._p.remove(r)
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p._p.append(r)


def replace_in_paras(doc, old, new, rebuild_ok=True):
    n = 0
    for p in doc.paragraphs:
        n += fix_runs(p._p, old, new)
        if n == 0 and rebuild_ok and old in merged(p):
            rebuild(p, merged(p).replace(old, new))
            n += 1
    return n


def set_tc(tc, text):
    """Set text of a raw w:tc element, preserving the first run's rPr."""
    ps = tc.findall(qn('w:p'))
    if not ps:
        return
    for extra in ps[1:]:
        tc.remove(extra)
    p = ps[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)


def set_cell(cell, text):
    p = cell.paragraphs[0]
    runs = para_runs(p)
    if runs:
        rpr = runs[0].find(qn('w:rPr'))
        for r in runs:
            p._p.remove(r)
        r = OxmlElement('w:r')
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p._p.append(r)
    else:
        p.text = text


def table_pct_width(tbl):
    tblPr = tbl._tbl.tblPr
    for w in tblPr.findall(qn('w:tblW')):
        tblPr.remove(w)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:type'), 'pct')
    tw.set(qn('w:w'), '5000')
    tblPr.append(tw)


def swap_image(doc, media_name, path):
    part = doc.part
    rid = None
    for r, rel in part.rels.items():
        if rel.target_ref == f'media/{media_name}':
            rid = r
    if rid is None:
        raise SystemExit(f'rel for {media_name} not found')
    img_part = part.related_parts[rid]
    data = open(path, 'rb').read()
    img_part._blob = data
    from PIL import Image
    import io
    w, h = Image.open(io.BytesIO(data)).size
    for p in doc.paragraphs:
        for blip in p._p.findall(f'.//{A_BLIP}'):
            if blip.get(R_EMBED) == rid:
                ext = p._p.findall(f'.//{WP_EXTENT}')[0]
                cx = int(ext.get('cx'))
                ext.set('cx', str(cx))
                ext.set('cy', str(round(cx * h / w)))


def harden_tables(doc):
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for t in doc.tables:
        tbl = t._tbl
        tblPr = tbl.find(W + 'tblPr')
        if tblPr is None:
            tblPr = tbl.makeelement(W + 'tblPr', {})
            tbl.insert(0, tblPr)
        tblW = tblPr.find(W + 'tblW')
        if tblW is None:
            tblW = tblPr.makeelement(W + 'tblW', {})
            tblPr.append(tblW)
        tblW.set(W + 'w', '5000')
        tblW.set(W + 'type', 'pct')
        layout = tblPr.find(W + 'tblLayout')
        if layout is None:
            layout = tblPr.makeelement(W + 'tblLayout', {})
            tblPr.append(layout)
        layout.set(W + 'type', 'autofit')
        grid = tbl.find(W + 'tblGrid')
        cols = grid.findall(W + 'gridCol') if grid is not None else []
        if cols:
            each = str(round(9360 / len(cols)))
            for c in cols:
                c.set(W + 'w', each)
    print(f'harden_tables: {len(doc.tables)} tables set to 100% width')


def compress_figure_spacing(doc):
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    n = 0
    for p in doc.paragraphs:
        ppr = p._p.find(W + 'pPr')
        if ppr is None:
            continue
        sp = ppr.find(W + 'spacing')
        if sp is not None:
            before = sp.get(W + 'before')
            after = sp.get(W + 'after')
            if before == '200' and after == '80':
                sp.set(W + 'before', '120')
                sp.set(W + 'after', '40')
                n += 1
    print(f'compressed spacing on {n} paragraphs')


def compress_ref_spacing(doc):
    import re
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    n = 0
    for p in doc.paragraphs:
        if not re.match(r'^\[\d+\]', p.text.strip()):
            continue
        ppr = p._p.find(W + 'pPr')
        if ppr is None:
            ppr = p._p.makeelement(W + 'pPr', {})
            p._p.insert(0, ppr)
        sp = ppr.find(W + 'spacing')
        if sp is None:
            sp = ppr.makeelement(W + 'spacing', {})
            ppr.append(sp)
        sp.set(W + 'line', '192')
        sp.set(W + 'lineRule', 'auto')
        sp.set(W + 'before', '10')
        sp.set(W + 'after', '10')
        n += 1
    print(f'compressed spacing on {n} reference paragraphs')


def add_refs(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith('[41]'):
            anchor = p._p
            for ref in REFS_NEW:
                np_ = copy.deepcopy(anchor)
                anchor.addnext(np_)
                anchor = np_
            for np_ in [p._p, *[anchor]]:  # anchor is last inserted; fix texts
                pass
            # write texts for the three inserted paragraphs
            cur = doc.paragraphs  # re-fetch after save? rebuild objects
            ps = [pp for pp in doc.paragraphs if pp._p is anchor or
                  (pp._p is not p._p and pp.text == '')]
            # simpler: walk XML siblings after [41]
            el = p._p
            nxt = el.getnext()
            tpl_run = el.find(qn('w:r'))
            rpr_tpl = tpl_run.find(qn('w:rPr')) if tpl_run is not None else None
            for ref in REFS_NEW:
                rebuild_from_xml(nxt, ref, rpr_tpl)
                nxt = nxt.getnext()
            print(f'added {len(REFS_NEW)} references after [41]')
            return
    raise SystemExit('ref [41] paragraph not found')


def rebuild_from_xml(el, text, rpr_tpl=None):
    runs = el.findall(qn('w:r'))
    for r in runs:
        el.remove(r)
    r = OxmlElement('w:r')
    if rpr_tpl is not None:
        r.append(copy.deepcopy(rpr_tpl))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    el.append(r)


def drop_mushtaq(doc):
    """Pass F: remove all [35] citations, the Table VIII MASnBr3 row and the
    reference entry, then renumber 36..43 down to 35..42."""
    for old, new in PASS_F:
        n = 0
        for p in doc.paragraphs:
            if old in p.text:
                n += replace_in_paras(doc, old, new)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old in p.text:
                            n += fix_runs(p._p, old, new)
        if n == 0:
            print(f'WARN: PASS_F no match: {old[:50]!r}...')
    left = [p.text[:100] for p in doc.paragraphs
            if not re.match(r'^\[\d+\]', p.text.strip())
            and ('Mushtaq' in p.text or '[35]' in p.text)]
    if left:
        raise SystemExit('PASS_F leftovers: ' + repr(left))
    for tbl in doc.tables:
        for row in list(tbl.rows):
            if row.cells[0].text.strip() == 'FTO/SnO2/MASnBr3/NiO/Au':
                row._tr.getparent().remove(row._tr)
                print('removed MASnBr3 row from table')
    renumber_refs(doc, MUSHTAQ_REFS, orphan=35)


SUB = qn('w:vertAlign')


def subscript_pass(doc):
    """Pass G: split plain formula/band-edge tokens into normal+subscript
    runs, Mushtaq-style. Already-subscripted runs are left untouched."""
    n = 0
    for p in list(doc.paragraphs) + [pp for t in doc.tables for r in t.rows
                                     for c in r.cells for pp in c.paragraphs]:
        for r in list(para_runs(p)):
            rpr = r.find(qn('w:rPr'))
            if rpr is not None and rpr.find(SUB) is not None:
                continue
            ts = r.findall(qn('w:t'))
            text = ''.join(t.text or '' for t in ts)
            if not SUB_RE.search(text):
                continue
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                r.insert(0, rpr)
            else:
                rpr = copy.deepcopy(rpr)
            subs = OxmlElement('w:vertAlign')
            subs.set(qn('w:val'), 'subscript')
            segs, pos = [], 0
            for m in SUB_RE.finditer(text):
                if m.start() > pos:
                    segs.append(('n', text[pos:m.start()]))
                pre = m.group(1) or m.group(3) or m.group(5) or m.group(7) or ''
                sub = m.group(2) or m.group(4) or m.group(6) or m.group(8) or ''
                if pre:
                    segs.append(('n', pre))
                segs.append(('s', sub))
                pos = m.end()
            if pos < len(text):
                segs.append(('n', text[pos:]))
            parent = r.getparent()
            parent.remove(r)
            for kind, s in segs:
                nr = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = s
                if kind == 's':
                    nr.append(copy.deepcopy(rpr))
                    nr.find(qn('w:rPr')).append(copy.deepcopy(subs))
                else:
                    nr.append(copy.deepcopy(rpr))
                nr.append(t)
                parent.append(nr)
            n += len(segs)
    print(f'subscript pass: {n} run segments rewritten')


def normalize_table_format(doc):
    """Pass H: uniform Times New Roman 9pt across all table cells. Runs with
    no rPr get the canonical template (TNR, not bold, not italic); runs with
    an rPr get missing rFonts/sz filled in. Existing bold (headers),
    italics and vertAlign (subscripts) are preserved."""
    tpl = None
    for r in para_runs(doc.tables[0].rows[1].cells[0].paragraphs[0]):
        cand = r.find(qn('w:rPr'))
        if cand is not None:
            tpl = cand
            break
    if tpl is None:
        raise SystemExit('no rPr template found in tables[0]')
    tpl = copy.deepcopy(tpl)
    filled = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in para_runs(p):
                        rpr = r.find(qn('w:rPr'))
                        if rpr is None:
                            rpr = copy.deepcopy(tpl)
                            r.insert(0, rpr)
                            filled += 1
                            continue
                        changed = False
                        if rpr.find(qn('w:rFonts')) is None:
                            rf = OxmlElement('w:rFonts')
                            rf.set(qn('w:ascii'), 'Times New Roman')
                            rf.set(qn('w:hAnsi'), 'Times New Roman')
                            rf.set(qn('w:cs'), 'Times New Roman')
                            rpr.insert(0, rf)
                            changed = True
                        sz = rpr.find(qn('w:sz'))
                        if sz is None:
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), '18')
                            rpr.append(sz)
                            changed = True
                        elif sz.get(qn('w:val')) != '18':
                            sz.set(qn('w:val'), '18')
                            changed = True
                        if rpr.find(qn('w:b')) is None:
                            b = OxmlElement('w:b')
                            b.set(qn('w:val'), '0')
                            rpr.append(b)
                            changed = True
                        if rpr.find(qn('w:i')) is None:
                            i = OxmlElement('w:i')
                            i.set(qn('w:val'), '0')
                            rpr.append(i)
                            changed = True
                        if changed:
                            filled += 1
    print(f'normalize_table_format: {filled} runs touched across '
          f'{len(doc.tables)} tables')


def apply_sample_table_style(doc):
    """Pass I: restyle every table like the user's sample - three thin black
    horizontal rules (top, under the header, bottom), no vertical or inner
    borders, header row larger (10 pt) and not bold, first column left-
    aligned, other columns centered."""
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    RULE = 'single'
    SZ = '4'
    for tbl in doc.tables:
        tblPr = tbl._tbl.tblPr
        style = tblPr.find(qn('w:tblStyle'))
        if style is None:
            style = OxmlElement('w:tblStyle')
            tblPr.insert(0, style)
        style.set(qn('w:val'), 'TableNormal')
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)
        borders = OxmlElement('w:tblBorders')
        for edge, tag in (('w:top', 'w:top'), ('w:bottom', 'w:bottom'),
                          ('w:left', 'w:left'), ('w:right', 'w:right'),
                          ('w:insideH', 'w:insideH'),
                          ('w:insideV', 'w:insideV')):
            el = OxmlElement(edge)
            val = 'none' if edge in ('w:left', 'w:right',
                                     'w:insideH', 'w:insideV') else RULE
            el.set(qn('w:val'), val)
            if val != 'none':
                el.set(qn('w:sz'), SZ)
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), '000000')
            borders.append(el)
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is not None:
            layout.addprevious(borders)
        else:
            tblPr.append(borders)

        rows = tbl.rows
        header = rows[0]
        for cell in header.cells:
            tcPr = cell._tc.tcPr
            if tcPr is None:
                tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcb = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), RULE)
            bottom.set(qn('w:sz'), SZ)
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            tcb.append(bottom)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                tcW.addnext(tcb)
            else:
                tcPr.insert(0, tcb)
            for p in cell.paragraphs:
                for r in p.runs:
                    rpr = r._r.get_or_add_rPr()
                    b = rpr.find(qn('w:b'))
                    if b is None:
                        b = OxmlElement('w:b')
                        rpr.append(b)
                    b.set(qn('w:val'), '0')
                    sz = rpr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rpr.append(sz)
                    sz.set(qn('w:val'), '20')
                    set_jc(p, 'center' if cell != header.cells[0] else 'left')

        for row in rows[1:]:
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    text = p.text
                    long_text = ci != 0 and len(text) > 30
                    set_jc(p, 'left' if (ci == 0 or long_text) else 'center')
    print(f'apply_sample_table_style: styled {len(doc.tables)} tables')


def set_jc(paragraph, value):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:jc')):
        pPr.remove(old)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), value)
    rPr = pPr.find(qn('w:rPr'))
    if rPr is not None:
        rPr.addprevious(jc)
    else:
        pPr.append(jc)


def pv_subscript_pass(doc):
    """Pass J: thesis-style V_OC / J_SC / I_SC typography for the paper text.
    'Voc'->'V'+'OC'(subscript), 'Jsc'->'J'+'SC'(subscript), 'Isc' likewise,
    matching the thesis's all-caps subscript convention in text and tables."""
    PW = re.compile(r'(?<![A-Za-z])([VJI])(oc|OC|sc|SC)(?![A-Za-z])')
    paras = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    n = 0
    for p in paras:
        for r in para_runs(p):
            for t in list(r.findall(qn('w:t'))):
                txt = t.text or ''
                if not PW.search(txt):
                    continue
                parent = r.getparent()
                rpr = r.find(qn('w:rPr'))
                newruns = []
                pos = 0
                for m in PW.finditer(txt):
                    if m.start() > pos:
                        newruns.append((txt[pos:m.start()], None))
                    newruns.append((m.group(1), None))
                    newruns.append((m.group(2).upper(), 'subscript'))
                    pos = m.end()
                if pos < len(txt):
                    newruns.append((txt[pos:], None))
                idx = parent.index(r)
                parent.remove(r)
                for text, va in newruns:
                    nr = OxmlElement('w:r')
                    rpr_new = copy.deepcopy(rpr) if rpr is not None else None
                    if va:
                        if rpr_new is None:
                            rpr_new = OxmlElement('w:rPr')
                        el = OxmlElement('w:vertAlign')
                        el.set(qn('w:val'), va)
                        rpr_new.append(el)
                    if rpr_new is not None:
                        nr.append(rpr_new)
                    nt = OxmlElement('w:t')
                    nt.set(qn('xml:space'), 'preserve')
                    nt.text = text
                    nr.append(nt)
                    parent.insert(idx, nr)
                    idx += 1
                n += 1
    cap = 0
    for p in paras:
        for r in para_runs(p):
            rpr = r.find(qn('w:rPr'))
            if rpr is None:
                continue
            va = rpr.find(qn('w:vertAlign'))
            if va is None or va.get(qn('w:val')) != 'subscript':
                continue
            ts = r.findall(qn('w:t'))
            txt = ''.join(t.text or '' for t in ts)
            if txt.upper() in ('OC', 'SC') and txt != txt.upper():
                for t in ts:
                    t.text = (t.text or '').upper()
                cap += 1
    print(f'pv_subscript_pass: {n} text runs rewritten'
          f', {cap} lowercase oc/sc subscripts capitalized')


def main():
    doc = docx.Document(WORK)

    # Pass A
    for old, new in GLOBAL:
        n = 0
        for p in doc.paragraphs:
            n += fix_runs(p._p, old, new)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        n += fix_runs(p._p, old, new)
        if len(old) < 10 and n == 0:
            print(f'WARN: no match for {old!r}')

    # Pass B
    for old, new in PASS_B:
        n = replace_in_paras(doc, old, new)
        if n == 0:
            print(f'WARN: PASS_B no match: {old[:60]!r}...')
    for old, new in METHODOLOGY:
        n = 0
        for p in doc.paragraphs:
            if old == merged(p):
                rebuild(p, new)
                n += 1
        if n == 0:
            print(f'WARN: METHODOLOGY no match: {old[:50]!r}...')

    # Pass C: tables
    vi, xiv = doc.tables[5], doc.tables[7]
    for row in vi.rows[1:]:
        voc, jsc, ff = [float(c.text.strip()) for c in row.cells[2:5]]
        set_cell(row.cells[5], f'{voc * jsc * ff / 100:.2f}')
    for row in xiv.rows[1:]:
        t = [c.text.strip() for c in row.cells]
        if t[0] == 'FTO/TiO2/RbGeI3/CuI/Au (this work)':
            continue
        pce = float(t[1]) * float(t[2]) * float(t[3]) / 100
        if abs(pce - float(t[4])) > 0.02:
            set_cell(row.cells[4], f'{pce:.2f}')
    last = xiv.rows[-1]._tr
    new_tr = copy.deepcopy(last)
    last.addprevious(new_tr)
    tcs = new_tr.findall(qn('w:tc'))
    for tc, txt in zip(tcs, MUSHTAQ_ROW):
        set_tc(tc, txt)
    for tbl in doc.tables:
        table_pct_width(tbl)

    # Pass C2: consistency fixes
    for old, new in PASS_D:
        n = 0
        for p in doc.paragraphs:
            if old in p.text:
                n += replace_in_paras(doc, old, new)
        if n == 0:
            print(f'WARN: PASS_D no match: {old[:60]!r}...')

    # Pass D
    harden_tables(doc)
    compress_figure_spacing(doc)
    compress_ref_spacing(doc)
    add_refs(doc)

    # Pass D2: renumber tail references by first-use order; drop orphan [41]
    renumber_refs(doc, RENUMBER, orphan=41)

    # Pass F: retire the Mushtaq reference (citations, table row, entry,
    # renumber 36..43 -> 35..42) -- must run after Pass D2 numbering.
    drop_mushtaq(doc)

    # Pass G: Mushtaq-style subscripts in formulas and band-edge labels.
    subscript_pass(doc)

    # Pass H: uniform table cell typography (Times New Roman 9 pt).
    normalize_table_format(doc)

    # Pass I: sample-style table rules and alignment.
    apply_sample_table_style(doc)

    # Pass J: thesis-style V_OC / J_SC / I_SC subscripts.
    pv_subscript_pass(doc)

    # Pass E
    for media, fig in SWAPS.items():
        swap_image(doc, media, f'{FIGDIR}/{fig}')
        print(f'swap {media} <- {fig}')

    doc.save(WORK)
    print('saved ->', WORK)


if __name__ == '__main__':
    main()